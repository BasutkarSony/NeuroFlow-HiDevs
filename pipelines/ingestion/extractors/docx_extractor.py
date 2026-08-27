from io import BytesIO

from docx import Document

from . import ExtractedPage


def extract_docx(file_bytes: bytes) -> list[ExtractedPage]:
    document = Document(BytesIO(file_bytes))
    pages: list[ExtractedPage] = []
    page_number = 1

    for paragraph in document.paragraphs:
        content = paragraph.text.strip()

        if not content:
            continue

        metadata = {}

        if paragraph.style and paragraph.style.name.startswith("Heading"):
            level = paragraph.style.name.replace("Heading", "").strip()

            if level.isdigit():
                metadata = {
                    "level": f"h{level}",
                    "section": content,
                }

        pages.append(
            ExtractedPage(
                page_number=page_number,
                content=content,
                content_type="text",
                metadata=metadata,
            )
        )

        page_number += 1

    for table_index, table in enumerate(document.tables, start=1):
        rows = []

        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
            ]
            rows.append("| " + " | ".join(cells) + " |")

        if rows:
            column_count = len(table.rows[0].cells)

            separator = (
                "| "
                + " | ".join(["---"] * column_count)
                + " |"
            )

            markdown = (
                rows[0]
                + "\n"
                + separator
                + "\n"
                + "\n".join(rows[1:])
            )

            pages.append(
                ExtractedPage(
                    page_number=page_number,
                    content=markdown,
                    content_type="table",
                    metadata={
                        "table_index": table_index,
                    },
                )
            )

            page_number += 1

    for section_index, section in enumerate(
        document.sections,
        start=1,
    ):
        for paragraph in section.header.paragraphs:
            content = paragraph.text.strip()

            if content:
                pages.append(
                    ExtractedPage(
                        page_number=page_number,
                        content=content,
                        content_type="text",
                        metadata={
                            "source": "header",
                            "section_index": section_index,
                        },
                    )
                )
                page_number += 1

    return pages